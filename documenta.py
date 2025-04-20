import json
from litellm import completion
import pandas as pd
import io
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import date, datetime
from chunkipy import TextChunker

# Funções de definição dos Prompts para a medida e fontes dos dados

def defined_prompt():
    """Retorna o prompt para a documentação do relatório do Power BI."""
    prompt_relatorio = """1 - Você é um documentador especializado em Power BI. Sua função é criar documentações claras e detalhadas para os relatórios, tabelas, medidas e fontes de dados em Power BI. Para cada item, você deve incluir uma descrição compreensiva que ajude os usuários a entenderem sua finalidade e uso no contexto do relatório. Utilize uma linguagem técnica e precisa, mas acessível para usuários com diferentes níveis de conhecimento em Power BI.
2 - Fazer a documentação no formato JSON.
3 - Você deverá dividir em diferentes outputs de acordo com a entrada do usuário, separando em: info_paineis, tabelas, medidas e fonte_de_dados.
4 - Na parte das medidas, você deverá fazer em blocos, das que estiverem sendo solicitadas, mas como continuação do JSON e ao final de todas fechar o JSON igual no exemplo.
5 - Retorne apenas o JSON, sem o ```JSON no inicio e o ``` no final
6 - O JSON deve ser retornado com aspas duplas, não simples.

Instruções Específicas:

Relatórios:
- Título do Relatório
- Descrição do objetivo do relatório
- Principais KPIs e métricas apresentadas
- Público-alvo do relatório
- Exemplos de uso

Formato de Documentação:

Tabelas do Relatório
Nome da Tabela | Descrição da Tabela

Medidas do Relatório
Nome da Medida | Descrição da Medida | Medida DAX

Fontes de Dados
Nome da Fonte de Dados | Descrição da Fonte | Tabelas Contidas no M | Nome da Tabela

Exemplo do JSON:
{
    "Relatorio": {
        "Titulo": "Análise de Vendas Mensais",
        "Descricao": "Este relatório fornece uma visão detalhada das vendas mensais por região e produto. Os principais KPIs incluem receita total, unidades vendidas e margem de lucro. O relatório é destinado aos gerentes de vendas regionais e é atualizado semanalmente para refletir os dados mais recentes.",
        "Principais_KPIs_e_Metricas": [
            "Receita Total",
            "Unidades Vendidas",
            "Margem de Lucro"
        ],
        "Publico_Alvo": "Gerentes de Vendas Regionais",
        "Exemplos_de_Uso": [
            "Identificação de tendências de vendas por região",
            "Comparação de desempenho de produtos"
        ]
    },
    "Tabelas_do_Relatorio": [
        {
            "Nome": "Vendas",
            "Descricao": "Tabela que armazena dados de vendas, incluindo ID do produto, quantidade vendida, preço e data da venda."
        },
        {
            "Nome": "Produtos",
            "Descricao": "Tabela que contém informações detalhadas dos produtos, como nome, categoria e preço unitário."
        }
    ],
    "Medidas_do_Relatorio": [
        {
            "Nome": "Receita Total",
            "Descricao": "Calcula a receita total das vendas somando o preço de venda multiplicado pela quantidade vendida."
        },
        {
            "Nome": "Margem de Lucro",
            "Descricao": "Calcula a margem de lucro subtraindo o custo do preço de venda."
        }
    ],
    "Fontes_de_Dados": [
        {
            "Nome": "SQL Server - Vendas",
            "Descricao": "Base de dados contendo todas as transações de vendas da empresa.",
            "Tabelas_Contidas_no_M": [
                "Vendas",
                "Produtos"
            ],
            "NomeTabela": "Vendas"
        },
        {
            "Nome": "Excel - Orçamento",
            "Descricao": "Planilha contendo dados de orçamento anual por departamento.",
            "Tabelas_Contidas_no_M": [
                "Orçamento"
            ],
            "NomeTabela": "Orçamento"
        }
    ]
}

Abaixo estão dados do relatório do Power BI a ser documentado:"""

    return prompt_relatorio

def defined_prompt_medidas():
    """Retorna o prompt para a documentação do relatório do Power BI."""
    prompt_relatorio = """1 - Você é um documentador especializado em Power BI. Sua função é criar documentações claras e detalhadas para os relatórios, tabelas, medidas e fontes de dados em Power BI. Para cada item, você deve incluir uma descrição compreensiva que ajude os usuários a entenderem sua finalidade e uso no contexto do relatório. Utilize uma linguagem técnica e precisa, mas acessível para usuários com diferentes níveis de conhecimento em Power BI.
2 - Fazer a documentação no formato JSON.
3 - Você deverá dividir em diferentes outputs de acordo com a entrada do usuário, separando em: info_paineis, tabelas e medidas.
4 - Na parte das medidas, você deverá fazer em blocos, das que estiverem sendo solicitadas, mas como continuação do JSON e ao final de todas fechar o JSON igual no exemplo.
5 - Retorne apenas o JSON, sem o ```JSON no inicio e o ``` no final
6 - Retornar somente um texto sem markdown, apenas o JSON como texto puro.
7 - O JSON deve ser retornado com aspas duplas, não simples.
8 - Importante levar em conta que as medidas do relatório podem ser enviadas por partes de acordo com o limite de tokens do modelo.

Instruções Específicas:

Relatórios:
- Título do Relatório
- Descrição do objetivo do relatório
- Principais KPIs e métricas apresentadas
- Público-alvo do relatório
- Exemplos de uso

Formato de Documentação:

Tabelas do Relatório
Nome da Tabela | Descrição da Tabela

Medidas do Relatório
Nome da Medida | Descrição da Medida | Medida DAX

Exemplo do JSON:
{
    "Relatorio": {
        "Titulo": "Análise de Vendas Mensais",
        "Descricao": "Este relatório fornece uma visão detalhada das vendas mensais por região e produto. Os principais KPIs incluem receita total, unidades vendidas e margem de lucro. O relatório é destinado aos gerentes de vendas regionais e é atualizado semanalmente para refletir os dados mais recentes.",
        "Principais_KPIs_e_Metricas": [
            "Receita Total",
            "Unidades Vendidas",
            "Margem de Lucro"
        ],
        "Publico_Alvo": "Gerentes de Vendas Regionais",
        "Exemplos_de_Uso": [
            "Identificação de tendências de vendas por região",
            "Comparação de desempenho de produtos"
        ]
    },
    "Tabelas_do_Relatorio": [
        {
            "Nome": "Vendas",
            "Descricao": "Tabela que armazena dados de vendas, incluindo ID do produto, quantidade vendida, preço e data da venda."
        },
        {
            "Nome": "Produtos",
            "Descricao": "Tabela que contém informações detalhadas dos produtos, como nome, categoria e preço unitário."
        }
    ],
    "Medidas_do_Relatorio": [
        {
            "Nome": "Receita Total",
            "Descricao": "Calcula a receita total das vendas somando o preço de venda multiplicado pela quantidade vendida."
        },
        {
            "Nome": "Margem de Lucro",
            "Descricao": "Calcula a margem de lucro subtraindo o custo do preço de venda."
        }
    ],
}

Abaixo estão dados do relatório do Power BI a ser documentado:"""
    return prompt_relatorio

def defined_prompt_fontes():
    """Retorna o prompt para a documentação do relatório do Power BI."""
    prompt_relatorio = """1 - Você é um documentador especializado em Power BI. Sua função é criar documentações claras e detalhadas para os relatórios, tabelas, medidas e fontes de dados em Power BI. Para cada item, você deve incluir uma descrição compreensiva que ajude os usuários a entenderem sua finalidade e uso no contexto do relatório. Utilize uma linguagem técnica e precisa, mas acessível para usuários com diferentes níveis de conhecimento em Power BI.
2 - Fazer a documentação no formato JSON.
3 - Você deverá dividir em diferentes outputs de acordo com a entrada do usuário, separando em: info_paineis, tabelas e fonte_de_dados.
5 - Retorne apenas o JSON, sem o ```JSON no inicio e o ``` no final
6 - Retornar somente um texto sem markdown, apenas o JSON como texto puro.
7 - O JSON deve ser retornado com aspas duplas, não simples.
8 - Importante levar em conta que as fontes dos dados das tabelas do relatório podem ser enviadas por partes de acordo com o limite de tokens do modelo.


Instruções Específicas:

Relatórios:
- Título do Relatório
- Descrição do objetivo do relatório
- Principais KPIs e métricas apresentadas
- Público-alvo do relatório
- Exemplos de uso

Formato de Documentação:

Tabelas do Relatório
Nome da Tabela | Descrição da Tabela

Fontes de Dados
Nome da Fonte de Dados | Descrição da Fonte | Tabelas Contidas no M | Nome da Tabela

Exemplo do JSON:
{
    "Relatorio": {
        "Titulo": "Análise de Vendas Mensais",
        "Descricao": "Este relatório fornece uma visão detalhada das vendas mensais por região e produto. Os principais KPIs incluem receita total, unidades vendidas e margem de lucro. O relatório é destinado aos gerentes de vendas regionais e é atualizado semanalmente para refletir os dados mais recentes.",
        "Principais_KPIs_e_Metricas": [
            "Receita Total",
            "Unidades Vendidas",
            "Margem de Lucro"
        ],
        "Publico_Alvo": "Gerentes de Vendas Regionais",
        "Exemplos_de_Uso": [
            "Identificação de tendências de vendas por região",
            "Comparação de desempenho de produtos"
        ]
    },
    "Tabelas_do_Relatorio": [
        {
            "Nome": "Vendas",
            "Descricao": "Tabela que armazena dados de vendas, incluindo ID do produto, quantidade vendida, preço e data da venda."
        },
        {
            "Nome": "Produtos",
            "Descricao": "Tabela que contém informações detalhadas dos produtos, como nome, categoria e preço unitário."
        }
    ],
    "Fontes_de_Dados": [
        {
            "Nome": "SQL Server - Vendas",
            "Descricao": "Base de dados contendo todas as transações de vendas da empresa.",
            "Tabelas_Contidas_no_M": [
                "Vendas",
                "Produtos"
            ],
            "NomeTabela": "Vendas"
        },
        {
            "Nome": "Excel - Orçamento",
            "Descricao": "Planilha contendo dados de orçamento anual por departamento.",
            "Tabelas_Contidas_no_M": [
                "Orçamento"
            ],
            "NomeTabela": "Orçamento"
        }
    ]
}

Abaixo estão dados do relatório do Power BI a ser documentado:"""
    return prompt_relatorio

# Define a tag para fazer a quebra do texto
def split_by_tag(text):
    return [t for t in text.split("<tag>") if t != '' and ' ']

def client_chat_LiteLLM(modelo, messages, maxtokens=4096):    
    """Interage com qualquer modelo unsando LiteLLM para obter respostas.
       Mais informações em: https://docs.litellm.ai/docs/providers
    """    
    count = 1
    try:
        response = completion(
            model=modelo,
            temperature=0,
            max_tokens=maxtokens,
            messages=messages
        )
        
        model_response = response.choices[0].message.content
        
        #remove the ```json and ``` from the response
        model_response = model_response.replace('```json', '').replace('```', '').replace('```JSON', '')
        
        #save the response to a file for debugging wih timestamp
        #timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")                
        #with open('response_' + timestamp + '_' +  modelo.replace('/', '_') + '.json', 'w', encoding='utf-8') as f:        
        #    f.write(model_response)
        
        response_content = json.loads( model_response )

        count += 1
    except Exception as e:
        print(f"Erro ao chamar a API do modelo: {modelo}, corrigindo automaticamente executando o modelo gpt-4.1-nano da Open AI. {str(e)}. Tentativa:", count)

        response = completion(
            model='gpt-4.1-nano',
            temperature=0,
            max_tokens=maxtokens,
            messages=messages
        )
        response_content = json.loads( response.choices[0].message.content )
        count += 1
        print(f"Modelo gpt-4.1-nano da Open AI executado com sucesso.")
        if count > 10:
            print(f"Erro ao chamar a API {str(e)}")
            raise Exception(f"Erro ao chamar a API {str(e)}")
        
        
    return response_content

def Documenta(prompt, text, modelo, max_tokens=4096, max_tokens_saida=4096):
    """Gera a documentação do relatório em formato JSON."""
    
    messages = [
        {"role": "system", "content": "Você é um documentador especializado em relatórios do Power BI."},
        {"role": "user", "content": f"{prompt}\n<INICIO DADOS RELATORIO POWER BI>\n{text}\n<FIM DADOS RELATORIO POWER BI>"}
    ]
    
    print('Usando o modelo:', modelo, 'Máximo de tokens de saída:', max_tokens_saida)
    
    response = client_chat_LiteLLM(modelo, messages, max_tokens_saida)
        
    return response

def set_heading(doc, text, level=1):
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    run.bold = True
    run.font.size = Pt(14 + (1 - level) * 2)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

def add_bullet_list(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style='List Bullet')
        run = paragraph.add_run(item)
        run.font.size = Pt(11)

def set_column_width(column, width):
    for cell in column.cells:
        cell.width = width

def add_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tcBorders.append(border)
            tcPr.append(tcBorders)

def style_table_header(cell):
    """Apply style to the table header cell."""
    # Set background color
    cell_fill = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), '000080')  # Navy blue background
    cell_fill.append(shd)
    # Set font color to white
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)  # White text

def add_measure_table(doc, measures, measures_df):
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Nome'
    hdr_cells[1].text = 'Descrição'
    hdr_cells[2].text = 'Fórmula DAX'
    
    # Style header cells
    for cell in hdr_cells:
        style_table_header(cell)
    
    # Adjust column widths
    widths = [Inches(1.5), Inches(6.0), Inches(2.5)]  # Set appropriate widths for columns
    for i, width in enumerate(widths):
        set_column_width(table.columns[i], width)
    
    def add_measure_row(measure):
        measure_name = measure["Nome"]
        #expression = measures_df.loc[measures_df['NomeMedida'] == measure_name, 'ExpressaoMedida'].values[0]
        
        matching_rows = measures_df.loc[measures_df['NomeMedida'] == measure_name, 'ExpressaoMedida']
            
        if not matching_rows.empty:
            expression = matching_rows.values[0]
        else:
            # Handle the case where there is no matching measure
            expression = "No mesure found"  # or any default value        
                
        row_cells = table.add_row().cells
        row_cells[0].text = measure_name
        row_cells[1].text = measure["Descricao"]
        row_cells[2].text = expression

    if isinstance(measures, list):
        for measure in measures:
            add_measure_row(measure)
    else:
        for measure in measures:
            add_measure_row(measure)

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.style = 'Body Text'
    
    add_table_borders(table)

def add_report_tables(doc, response_tables):
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Tabela'
    hdr_cells[1].text = 'Descrição'
    
    # Style header cells
    for cell in hdr_cells:
        style_table_header(cell)
    
    # Adjust column widths
    widths = [Inches(2.0), Inches(5.0)]  # Set appropriate widths for columns
    for i, width in enumerate(widths):
        set_column_width(table.columns[i], width)

    def add_table_row(table_info):
        row_cells = table.add_row().cells
        row_cells[0].text = table_info["Nome"]
        row_cells[1].text = table_info["Descricao"]

    if isinstance(response_tables, list):
        for table_info in response_tables:
            add_table_row(table_info)
    elif 'Tabelas_do_Relatorio' in response_tables:
        for table_info in response_tables['Tabelas_do_Relatorio']:
            add_table_row(table_info)
    
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.style = 'Body Text'
    
    add_table_borders(table)

def add_data_sources_table(doc, response_source):
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Nome'
    hdr_cells[1].text = 'Descrição'
    hdr_cells[2].text = 'Tabelas contidas no M'
    hdr_cells[3].text = 'Código M'
    
    # Style header cells
    for cell in hdr_cells:
        style_table_header(cell)
    
    # Adjust column widths
    widths = [Inches(2.0), Inches(4.0), Inches(2.0), Inches(2.0)]  # Set appropriate widths for columns
    for i, width in enumerate(widths):
        set_column_width(table.columns[i], width)

    def add_source_row(source_info):
        row_cells = table.add_row().cells
        row_cells[0].text = source_info["Nome"]
        row_cells[1].text = source_info["Descricao"]
        row_cells[2].text = ", ".join(source_info["Tabelas_Contidas_no_M"])         
        row_cells[3].text = source_info.get("FonteDados", "N/A")
    
    if isinstance(response_source, list):
        for source_info in response_source:
            add_source_row(source_info)
    elif 'Fontes_de_Dados' in response_source:
        for source_info in response_source['Fontes_de_Dados']:
            add_source_row(source_info)
    
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.style = 'Body Text'
    
    add_table_borders(table)

def add_centered_title(doc, title, color=RGBColor(0, 0, 128)):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = color

def generate_promt_medidas(text):
    
    prompts = defined_prompt_medidas().strip()
    
    return f"{prompts}\n<INICIO DADOS RELATORIO POWER BI>\n{text}\n<FIM DADOS RELATORIO POWER BI>"

def generate_promt_fontes(text):
    
    prompts = defined_prompt_fontes().strip()
    
    return f"{prompts}\n<INICIO DADOS RELATORIO POWER BI>\n{text}\n<FIM DADOS RELATORIO POWER BI>"

from docx.shared import Inches

def add_colunas_table(doc, df_colunas):
    table = doc.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Tabela'
    hdr_cells[1].text = 'Coluna'
    hdr_cells[2].text = 'Tipo'
    hdr_cells[3].text = 'Calculada ou Dados'
    hdr_cells[4].text = 'Expressão'
    
    # Style header cells
    for cell in hdr_cells:
        style_table_header(cell)
    
    # Adjust column widths
    widths = [Inches(2.0), Inches(2.0), Inches(2.0), Inches(1.5), Inches(3.0)]  # Set appropriate widths for columns
    for i, width in enumerate(widths):
        set_column_width(table.columns[i], width)

    # Add rows from dataframe
    for _, row in df_colunas.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['NomeTabela'])
        row_cells[1].text = str(row['NomeColuna'])
        row_cells[2].text = str(row['TipoDadoColuna'])
        row_cells[3].text = str(row['TipoColuna'])
        row_cells[4].text = str(row['ExpressaoColuna'])
    
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.style = 'Body Text'
    
    add_table_borders(table)

def add_relationamentos_table(doc, df_relacionamentos):
    # Criar a tabela com 4 colunas
    table = doc.add_table(rows=1, cols=4)
    
    # Definir os cabeçalhos da tabela
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'De tabela'
    hdr_cells[1].text = 'De coluna'
    hdr_cells[2].text = 'Para tabela'
    hdr_cells[3].text = 'Para coluna'
    
    # Estilizar as células do cabeçalho (se necessário)
    for cell in hdr_cells:
        style_table_header(cell)
    
    # Ajustar a largura das colunas
    widths = [Inches(2.0), Inches(2.0), Inches(2.0), Inches(2.0)]  # Defina a largura apropriada para cada coluna
    for i, width in enumerate(widths):
        set_column_width(table.columns[i], width)
    
    # Adicionar as linhas com os dados do dataframe
    for _, row in df_relacionamentos.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row['FromTable']
        row_cells[1].text = row['FromColumn']
        row_cells[2].text = row['ToTable']
        row_cells[3].text = row['ToColumn']
    
    # Estilizar o corpo da tabela
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.style = 'Body Text'
    
    # Adicionar bordas à tabela (se necessário)
    add_table_borders(table)

def generate_docx(response_info, response_tables, response_measures, response_source, measures_df, df_relationships, df_colunas):
    """Gera um documento Word com a documentação do relatório."""
    doc = Document()
    
    # Add Logo
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run()
    run.add_picture(r'images/Logo.png', width=Inches(1.0))
    
    # Add centered title
    add_centered_title(doc, "AutoDoc 2025 - Documentador de Power BI")
        
    # Title and Description
    set_heading(doc, 'Relatório:', level=1)
    doc.add_paragraph(response_info["Titulo"], style='Body Text')

    # Data
    set_heading(doc, 'Data:', level=1)
    today = date.today()
    doc.add_paragraph(today.strftime("%d/%m/%Y %H:%M:%S"), style='Body Text')
    
    set_heading(doc, 'Descrição:', level=1)
    doc.add_paragraph(response_info["Descricao"], style='Body Text')
    
    # Main KPIs and Metrics
    set_heading(doc, 'Principais KPIs e Métricas:', level=1)
    add_bullet_list(doc, response_info["Principais_KPIs_e_Metricas"])
    
    # Target Audience
    set_heading(doc, 'Público alvo:', level=1)
    doc.add_paragraph(response_info["Publico_Alvo"], style='Body Text')
    
    # Use Cases
    set_heading(doc, 'Exemplos de uso:', level=1)
    add_bullet_list(doc, response_info["Exemplos_de_Uso"])
    
    # Report Tables
    set_heading(doc, 'Tabelas', level=1)
    add_report_tables(doc, response_tables)
    
    # Report Measures
    set_heading(doc, 'Medidas', level=1)
    add_measure_table(doc, response_measures, measures_df)
    
    # Data Sources
    set_heading(doc, 'Fonte de dados', level=1)
    add_data_sources_table(doc, response_source)

    # Columns
    set_heading(doc, 'Colunas', level=1)
    add_colunas_table(doc, df_colunas)

    # Relationships
    if df_relationships is not None:
        set_heading(doc, 'Relacionamentos', level=1)
        add_relationamentos_table(doc, df_relationships)

    return doc

def generate_excel(response_info, response_tables, response_measures, response_source, measures_df, df_relationships, df_colunas):
    """Gera um arquivo Excel com a documentação do relatório."""
    buffer = io.BytesIO()
    
    with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
        all_info = []
        all_tables = []
        all_measures = []
        all_sources = []
                
        info = pd.DataFrame([response_info]).transpose()
        info.reset_index(inplace=True)
        info.columns = ['Informações do relatório', 'Dados']
        all_info.append(info)
        
        if isinstance(response_tables, list):
            tables = pd.DataFrame(response_tables)
        elif isinstance(response_tables, dict):
            tables = pd.DataFrame(response_tables)
        else:
            tables = pd.DataFrame()
            
        all_tables.append(tables)
        
        if isinstance(response_measures, list):
            measures = pd.DataFrame(response_measures)
        elif isinstance(response_measures, dict):
            measures = pd.DataFrame(response_measures)
        else:
            measures = pd.DataFrame()
            
        all_measures.append(measures)
        
        if isinstance(response_source, list):
            sources = pd.DataFrame(response_source)
        elif isinstance(response_source, dict):
            sources = pd.DataFrame(response_source)
        else:
            sources = pd.DataFrame()
            
        all_sources.append(sources)
    
        df_info = pd.concat(all_info, ignore_index=True)
        df_tabelas = pd.concat(all_tables, ignore_index=True)
        df_medidas = pd.concat(all_measures, ignore_index=True)
        df_fontes = pd.concat(all_sources, ignore_index=True)
        
        if isinstance(response_source, dict):
            df_medidas = pd.merge(df_medidas, measures_df,  left_on='Nome', right_on='Medida', how='left')
            df_medidas = df_medidas[['Medida', 'Descricao', 'expression']]
    
        df_info.to_excel(writer, sheet_name='info_painel', index=False)
        df_tabelas.to_excel(writer, sheet_name='tabelas', index=False) 
        df_medidas.to_excel(writer, sheet_name='medidas', index=False) 
        df_fontes.to_excel(writer, sheet_name='fonte_de_dados', index=False)

        if df_relationships is not None:
            df_relationships.to_excel(writer, sheet_name='relacionamentos', index=False)

        df_colunas.to_excel(writer, sheet_name='colunas', index=False)

            
    buffer.seek(0)
    return buffer

# Funçcão para preparar o relatório do Power BI para enviar para o modelo LLM por prompt

def text_to_document(df, df_relationships=None, max_tokens=4096):
    """Gera o texto para documentação baseado nos dados do DataFrame."""
    
    # Define o tamanho máximo de tokens para o modelo LLM
    pd.set_option('display.max_colwidth', None)
    
    # Faz a leitura dos dados do relatório do Power BI para a preparação para gerar o relatório
    tables_df = df[df['NomeTabela'].notnull() & df['FonteDados'].notnull()]
    tables_df = tables_df[['NomeTabela', 'FonteDados']].drop_duplicates().reset_index(drop=True)

    measures_df = df[df['NomeMedida'].notnull() & df['ExpressaoMedida'].notnull()]
    measures_df = measures_df[['NomeMedida', 'ExpressaoMedida']].drop_duplicates().reset_index(drop=True)

    df_colunas = df[['NomeTabela','NomeColuna', 'TipoDadoColuna', 'TipoColuna', 'ExpressaoColuna']]
    df_colunas = df_colunas[df_colunas['NomeTabela'] != 'Medidas']

    df_colunas['TipoColuna'] = df_colunas['TipoColuna'].replace('N/A', '')
    df_colunas['ExpressaoColuna'] = df_colunas['ExpressaoColuna'].replace('N/A', '')

    # filter the df_colunas not null
    df_colunas = df_colunas[df_colunas['NomeColuna'].notnull()]

    if not df.empty and 'ReportName' in df.columns:
        report_name = df['ReportName'].iloc[0]
    else:
        report_name = "PBIReport"

    if df_relationships is None:
        df_relationships = pd.DataFrame()

    # Prepara para enviar as medidas do relatório em partes por causa da limitação de tokens do modelo
    #monta um texto com o nome da medida e a expressao da medida    
    measures_df['NomeMedidaExpressao'] = '<tag> Nome da medida: ' + measures_df['NomeMedida'] + ' Expressão da medida: ' + measures_df['ExpressaoMedida']

    text_chunker_medidas = TextChunker(chunk_size=max_tokens, tokens=True, overlap_percent=0, split_strategies=[split_by_tag])
    chunks_medidas = text_chunker_medidas.chunk(measures_df['NomeMedidaExpressao'].to_string(index=False))

    # Prepara para enviar as fontes dos dados do relatório em partes por causa da limitação de tokens do modelo
    #monta um texto com o nome da tabela e fontededados
    tables_df['NomeTabelaFonteDados'] = '<tag> NomeTabela: ' + tables_df['NomeTabela'] + ' Fonte de Dados: ' + tables_df['FonteDados']

    text_chunker_fontes = TextChunker(chunk_size=max_tokens, tokens=True, overlap_percent=0, split_strategies=[split_by_tag])
    chunks_fontes = text_chunker_fontes.chunk(tables_df['NomeTabelaFonteDados'].to_string(index=False))

    # define uma lista para armazenar todos os textos para o modelo LLM
    document_texts_medidas = []
    document_texts_fontes = []

    # monta o texto baseados na medidas
    for chunk in chunks_medidas:
        document_texts_medidas.append(

            f"""
                  Relatório: {report_name}

                  Tabelas:
                  {tables_df['NomeTabela'].to_string(index=False)}

                  Medidas:
                  {chunk}
                  """
            )

    # monta o texto baseados nas fontes de dados
    for chunk in chunks_fontes:
        document_texts_fontes.append(

            f"""
                  Relatório: {report_name}

                  Fontes dos dados das tabelas:
                  {chunk}
                  """
            )

    # monta o texto final para o relatório
    document_text_all = f"""
        
    Relatório: {report_name}
    
    <Tabelas>
    {tables_df['NomeTabela'].to_string(index=False)}
    </Tabelas>
    
    <Medidas do Relatório>
    {measures_df['NomeMedidaExpressao'].to_string(index=False)}
    </Medidas do Relatório>
    
    <Fontes de dados>
    {tables_df['NomeTabelaFonteDados'].to_string(index=False)}
    </Fontes de dados>
    """ 
    
    return document_text_all, document_texts_medidas, document_texts_fontes, measures_df, tables_df, df_colunas